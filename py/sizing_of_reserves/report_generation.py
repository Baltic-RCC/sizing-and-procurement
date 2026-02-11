import logging
import math
import os
import webbrowser
from datetime import datetime
from dataclasses import dataclass

import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from fpdf import FPDF, XPos, YPos

from py.common.functions import get_file_path_by_folder
from py.common.time_functions import convert_datetime_to_string

logger = logging.getLogger(__name__)

PLUS_MINUS = u'\u00b1'
HEADER_FIGURE_X = 170
HEADER_FIGURE_Y = 1
HEADER_FIGURE_WIDTH = 40
HEADER_FIGURE_LOCATION = '../resources/rcc_logo.png'
ALL_DATA_KEYWORD = 'all'
POSITIVE_DATA_KEYWORD = '+'
NEGATIVE_DATA_KEYWORD = '-'
DET_DESCRIPTION_KEYWORD = 'Det.'
DETERMINISTIC_FIGURE_NAME = 'deterministic_results.png'
MC_DESCRIPTION_KEYWORD = 'MC'
MC_FIGURE_NAME = 'mc_results.png'
DATE_FORMAT_FOR_REPORT = "%d.%m.%Y"
TEXT_FONT_SIZE = 12
FONT_FAMILY = 'Helvetica'
HEADING_FONT_SIZE = 20
SUBHEADING_FONT_SIZE = 16
UNDER_SUBHEADING_FONT_SIZE = 14


class PDF(FPDF):

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)

    def header(self):
        """
        Initializes a custom header, will be used by super when add_page is called

        :return: None
        """
        file_abs_path = get_file_path_by_folder(HEADER_FIGURE_LOCATION)
        if os.path.isfile(file_abs_path):
            try:
                self.image(file_abs_path, x=HEADER_FIGURE_X, y=HEADER_FIGURE_Y, w=HEADER_FIGURE_WIDTH)
            except FileNotFoundError:
                logger.error(f"Logo was not found at {file_abs_path}")
        else:
            logger.error(f"Logo was not present at {file_abs_path}")
        self.ln(30)

    def footer(self):
        """
        Initializes a custom footer, will be used by super when add_page is called

        :return:
        """
        self.set_y(-15)
        self.set_font(FONT_FAMILY, "I", 10)
        self.cell(0, 10, f"{self.page_no()}/{{nb}}", align="C")


# class SizingWordDocument(Document):
#     pass

@dataclass
class SizingOutput:
    """
    Additional class to collect parameters regard to the report generation
    """
    region_list: list
    start_date: str | datetime = None
    end_date: str | datetime = None
    date_today: str | datetime = None
    region: str = None

    def __post_init__(self):
        """

        :return:
        """
        self.region = ", ".join(str(region_value) for region_value in self.region_list)
        if self.region.lower() == 'baltics':
            self.region = 'Baltic SOR'
        try:
            self.date_today = convert_datetime_to_string(self.date_today, DATE_FORMAT_FOR_REPORT)
        except ValueError:
            self.date_today = pd.Timestamp("today").strftime(DATE_FORMAT_FOR_REPORT)
        self.date_today = convert_datetime_to_string(self.date_today, DATE_FORMAT_FOR_REPORT)
        self.start_date = convert_datetime_to_string(self.start_date, DATE_FORMAT_FOR_REPORT)
        self.end_date = convert_datetime_to_string(self.end_date, DATE_FORMAT_FOR_REPORT)



def add_page_numbers_to_word_document(word_paragraph):
    """
    Generates document number for the Word document in style of {page_number}/{number_of_pages}

    :param word_paragraph: paragraph where to generate page numbers
    :return:
    """
    run = word_paragraph.add_run("")
    field_char_1 = OxmlElement('w:fldChar')
    field_char_1.set(qn('w:fldCharType'), 'begin')
    insert_text_1 = OxmlElement('w:instrText')
    insert_text_1.text = 'PAGE'
    field_char_2 = OxmlElement('w:fldChar')
    field_char_2.set(qn('w:fldCharType'), 'separate')
    field_char_3 = OxmlElement('w:fldChar')
    field_char_3.set(qn('w:fldCharType'), 'end')
    run._r.append(field_char_1)
    run._r.append(insert_text_1)
    run._r.append(field_char_2)
    run._r.append(field_char_3)
    run = word_paragraph.add_run("/")
    field_char_1 = OxmlElement('w:fldChar')
    field_char_1.set(qn('w:fldCharType'), 'begin')
    insert_text_1 = OxmlElement('w:instrText')
    insert_text_1.text = 'NUMPAGES'
    field_char_2 = OxmlElement('w:fldChar')
    field_char_2.set(qn('w:fldCharType'), 'separate')
    field_char_3 = OxmlElement('w:fldChar')
    field_char_3.set(qn('w:fldCharType'), 'end')
    run._r.append(field_char_1)
    run._r.append(insert_text_1)
    run._r.append(field_char_2)
    run._r.append(field_char_3)


def disable_autofit(input_table):
    """
    Disables autofitting for the tables (if needed)

    :param input_table: python-docx table instance
    :return: None
    """
    table_value = input_table._tbl
    table_pr = table_value.tblPr
    table_layout = OxmlElement('w:tblLayout')
    table_layout.set(qn('w:type'), 'fixed')
    table_pr.append(table_layout)


def generate_report_word(heading_string: str,
                        methodologies: dict,
                        summaries: dict,
                        region_list: list,
                        tables: dict,
                        images: dict,
                        references: dict = None,
                        file_name: str = None,
                        start_date: str | datetime = None,
                        end_date: str | datetime = None,
                        date_today: str | datetime = None):
    """
    Generates a word report from this analysis
    (Taken from https://towardsdtatascience.com/how-to-create-a-pdf-report-fro-your-data-analysis-in-python)

    :param end_date: end time moment for the calculation
    :param start_date: start time moment for the calculation
    :param references: list of references
    :param methodologies: Methodology of the analysis process
    :param heading_string: The title of the report
    :param summaries: Main results with comments
    :param region_list: List of regions where analysis was performed
    :param tables: dictionary of tables (table caption: pandas.dataframe as table)
    :param images: dictionary of image addresses (image caption: image location)
    :param file_name: name of the file where to save the report
    :param date_today:
    """
    word_doc = Document()

    word_section = word_doc.sections[0]
    page_width = word_section.page_width
    left_margin = word_section.left_margin
    right_margin = word_section.right_margin
    usable_width = page_width - left_margin -right_margin

    word_header = word_section.header
    header_paragraph = word_header.paragraphs[0]
    file_abs_path = get_file_path_by_folder(HEADER_FIGURE_LOCATION)
    run = header_paragraph.add_run()
    run.add_picture(file_abs_path)
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    word_footer = word_section.footer
    footer_paragraph = word_footer.paragraphs[0]
    footer_paragraph.clear()
    add_page_numbers_to_word_document(footer_paragraph)
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    word_doc.add_heading(text=heading_string, level=0)
    summary_table = word_doc.add_table(rows=4, cols=2)
    parameters = SizingOutput(region_list=region_list, start_date=start_date, end_date=end_date, date_today=date_today)

    summary_table.rows[0].cells[0].text = 'Date:'
    summary_table.rows[0].cells[1].text = parameters.date_today
    summary_table.rows[1].cells[0].text = 'Analysis is performed for '
    summary_table.rows[1].cells[1].text = parameters.region
    summary_table.rows[2].cells[0].text = 'From: '
    summary_table.rows[2].cells[1].text = parameters.start_date
    summary_table.rows[3].cells[0].text = 'To: '
    summary_table.rows[3].cells[1].text = parameters.end_date

    word_doc.add_heading(text="Methodology", level=1)

    for methodology in methodologies:
        word_doc.add_heading(methodology, level=2)
        methodology_paragraph = word_doc.add_paragraph(methodologies[methodology])
        methodology_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    word_doc.add_heading(text="Results", level=1)
    for summary in summaries:
        summary_paragraph = word_doc.add_paragraph(summaries[summary])
        summary_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    table_counter = 1
    for table_description in tables:
        table_heading = f"Table {table_counter}: {table_description}"
        table_heading_paragraph = word_doc.add_paragraph(text=table_heading)
        table_heading_paragraph.paragraph_format.space_before = Pt(12)
        table_heading_paragraph.paragraph_format.space_after = Pt(3)
        table_heading_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        table_counter += 1
        input_table = tables[table_description]
        word_table = word_doc.add_table(rows=1, cols=len(input_table.columns))
        word_table.style = 'Table Grid'
        header_cells = word_table.rows[0].cells
        for i, column_name in enumerate(input_table.columns):
            header_run = header_cells[i].paragraphs[0].add_run(column_name)
            header_run.bold = True
            # header_run.font.size = Pt(11)
            header_cells[i].paragraphs[0].alignment = 1
            # header_cells[i].text = str(column_name)

        for index, row in input_table.iterrows():
            row_cells = word_table.add_row().cells
            for i, item in enumerate(row):
                row_cell = row_cells[i]
                row_parameter = row_cell.paragraphs[0]
                row_run = row_parameter.add_run(str(item))
                # row_run.font.size = Pt(10)
                row_parameter.alignment = 0
                # row_cells[i].text = str(item)

    figure_counter = 1
    for image in images:
        pic_paragraph = word_doc.add_paragraph()
        run = pic_paragraph.add_run()
        run.add_picture(images[image])
        pic_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        pic_paragraph.paragraph_format.space_before = Pt(12)
        pic_paragraph.paragraph_format.space_after = Pt(3)
        figure_caption = f"Figure {figure_counter}: {image}"
        caption_paragraph = word_doc.add_paragraph(text=figure_caption)
        caption_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        # pic_paragraph.paragraph_format.space_before = Pt(12)
        pic_paragraph.paragraph_format.space_after = Pt(12)
        figure_counter = figure_counter + 1
    #
    if references is not None and len(references) > 0:
        word_doc.add_heading(text="References", level=1)
        word_references = word_doc.add_table(rows=len(references), cols=2)
        disable_autofit(word_references)
        number_column_width = Cm(2)
        text_column_width = usable_width - number_column_width
        for i, reference in enumerate(references):
            word_references.rows[i].cells[0].text = str(reference)
            word_references.rows[i].cells[0].width = number_column_width
            word_references.rows[i].cells[1].text = references[reference]
            word_references.rows[i].cells[1].width = text_column_width
    if file_name is not None:
        word_doc.save(file_name)
        return None
    return word_doc


def generate_report_pdf(heading_string: str,
                        methodologies: dict,
                        summaries: dict,
                        region_list: list,
                        tables: dict,
                        images: dict,
                        references: dict = None,
                        file_name: str = None,
                        start_date: str | datetime = None,
                        end_date: str | datetime = None,
                        date_today: str | datetime = None):
    """
    Generates a pdf report from this analysis
    (Taken from https://towardsdtatascience.com/how-to-create-a-pdf-report-fro-your-data-analysis-in-python)

    :param end_date: end time moment for the calculation
    :param start_date: start time moment for the calculation
    :param references: list of references
    :param methodologies: Methodology of the analysis process
    :param heading_string: The title of the report
    :param summaries: Main results with comments
    :param region_list: List of regions where analysis was performed
    :param tables: dictionary of tables (table caption: pandas.dataframe as table)
    :param images: dictionary of image addresses (image caption: image location)
    :param file_name: name of the file where to save the report
    :param date_today:
    """
    cell_height = 8
    line_height = TEXT_FONT_SIZE
    margin_between_paragraphs = 6
    parameters = SizingOutput(region_list=region_list, start_date=start_date, end_date=end_date, date_today=date_today)

    pdf_file = PDF("P", "mm", "A4")
    pdf_file.set_margins(left=10, top=10)
    pdf_file.set_text_color(r=0, g=0, b=0)
    pdf_file.add_page()
    pdf_file.set_font(FONT_FAMILY, 'B', HEADING_FONT_SIZE)
    pdf_file.multi_cell(w=0, h=1 * cell_height, text=heading_string, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf_file.set_font(FONT_FAMILY, '', TEXT_FONT_SIZE)
    pdf_file.cell(w=30, h=cell_height, text='Date:', new_x=XPos.RIGHT, new_y=YPos.TOP)
    pdf_file.cell(w=30, h=cell_height, text=parameters.date_today, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf_file.cell(w=60, h=cell_height, text='Analysis is performed for ', new_x=XPos.RIGHT, new_y=YPos.TOP)
    pdf_file.cell(w=60, h=cell_height, text=parameters.region, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf_file.cell(w=60, h=cell_height, text='From: ', new_x=XPos.RIGHT, new_y=YPos.TOP)
    pdf_file.cell(w=60, h=cell_height, text=parameters.start_date, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf_file.cell(w=60, h=cell_height, text='To:', new_x=XPos.RIGHT, new_y=YPos.TOP)
    pdf_file.cell(w=60, h=cell_height, text=parameters.end_date, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf_file.set_font(FONT_FAMILY, 'B', SUBHEADING_FONT_SIZE)
    pdf_file.cell(w=0, h=2 * cell_height, text="Methodology", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf_file.set_font(FONT_FAMILY, '', TEXT_FONT_SIZE)
    for methodology in methodologies:
        pdf_file.set_font(FONT_FAMILY, 'I', UNDER_SUBHEADING_FONT_SIZE)
        pdf_file.cell(w=0, h=2 * cell_height, text=methodology, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf_file.set_font(FONT_FAMILY, '', TEXT_FONT_SIZE)
        pdf_file.multi_cell(w=0, h=5, text=methodologies[methodology])
        pdf_file.ln(margin_between_paragraphs)

    pdf_file.set_font(FONT_FAMILY, 'B', SUBHEADING_FONT_SIZE)
    pdf_file.cell(w=0, h=2 * cell_height, text="Results", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf_file.set_font(FONT_FAMILY, '', TEXT_FONT_SIZE)
    for summary in summaries:
        pdf_file.multi_cell(w=0, h=5, text=summaries[summary])
        pdf_file.ln(margin_between_paragraphs)

    table_counter = 1
    for table_description in tables:
        pdf_file.set_font(FONT_FAMILY, '', TEXT_FONT_SIZE)
        table_heading = f"Table {table_counter}: {table_description}"
        pdf_file.multi_cell(w=0, h=5, text=table_heading)
        pdf_file.ln(margin_between_paragraphs)
        table_counter += 1
        table = tables[table_description]

        cell_width = (210 - 10 - 10) / (len(table.columns) + 1)

        number_of_lines = 1
        for column_name in table.columns:
            new_number_lines = math.ceil(pdf_file.get_string_width(str(column_name)) / cell_width)
            number_of_lines = max(new_number_lines, number_of_lines)
        pdf_file.multi_cell(w=cell_width,
                            h=line_height * number_of_lines * 1,
                            text="",
                            align="C",
                            border="B",
                            new_x="RIGHT",
                            new_y="TOP",
                            max_line_height=line_height)
        for column_name in table.columns:
            pdf_file.multi_cell(w=cell_width,
                                h=line_height * number_of_lines * 1,
                                text=str(column_name),
                                align="C",
                                border="B",
                                new_x="RIGHT",
                                new_y="TOP",
                                max_line_height=line_height)
        pdf_file.ln(line_height * 1 * number_of_lines)

        for index, row in table.iterrows():
            number_of_lines = 1
            new_number_lines = math.ceil(pdf_file.get_string_width(str(index)) / cell_width)
            number_of_lines = max(new_number_lines, number_of_lines)
            for i in range(len(table.columns)):
                new_number_lines = math.ceil(pdf_file.get_string_width(str(row.iloc[i])) / cell_width)
                number_of_lines = max(new_number_lines, number_of_lines)
            pdf_file.multi_cell(w=cell_width,
                                h=line_height * number_of_lines * 1,
                                text=str(index),
                                align="C",
                                border="B",
                                new_x="RIGHT",
                                new_y="TOP",
                                max_line_height=line_height)
            for i in range(len(table.columns)):
                pdf_file.multi_cell(w=cell_width,
                                    h=line_height * number_of_lines * 1,
                                    text=str(row.iloc[i]),
                                    align="C",
                                    border="B",
                                    new_x="RIGHT",
                                    new_y="TOP",
                                    max_line_height=line_height)
            pdf_file.ln(line_height * 1 * number_of_lines)
        pdf_file.ln(cell_height)
    pdf_file.ln(cell_height)

    figure_counter = 1
    for image in images:
        # if figure_counter > 1 & figure_counter % 2 == 1:
        #    pdf_file.add_page()
        pdf_file.image(images[image], w=pdf_file.epw)
        figure_caption = f"Figure {figure_counter}: {image}"
        pdf_file.ln(cell_height)
        pdf_file.set_font(FONT_FAMILY, '', TEXT_FONT_SIZE)
        pdf_file.multi_cell(w=0, h=5, text=figure_caption)
        pdf_file.ln(cell_height)
        figure_counter += 1

    if references is not None and len(references) > 0:
        pdf_file.ln(cell_height)
        pdf_file.set_font(FONT_FAMILY, 'B', SUBHEADING_FONT_SIZE)
        pdf_file.cell(w=0, h=2 * cell_height, text="References", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf_file.set_font(FONT_FAMILY, '', TEXT_FONT_SIZE)
        for reference in references:
            pdf_file.cell(w=10, h=5, text=str(reference), new_x=XPos.RIGHT, new_y=YPos.TOP)
            pdf_file.multi_cell(w=180, h=5, text=references[reference], new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf_file.ln(cell_height)

    if file_name is not None:
        pdf_file.output(file_name)
        webbrowser.open(file_name)
        return None
    else:
        return pdf_file.output()
